
# -*- coding: utf-8 -*-
"""
Report Builder (additive):
- Do NOT duplicate existing "Revenue Graph" or "TOP PRODUCTS TABLE".
- If multiple years -> line chart with labels; if single year -> NO chart.
- Update existing Top Products/Segments table in place via PDF-first (official Nokia filings), then web-search.
- If not found, explicitly write "Not publicly available" in the table cell.
- Cleanly render headings, bullets, links; strip escaped markdown artifacts.
- Word tables: Table Grid style, header shading, column widths, center alignment.
Two entry points:
 build_full_report(company, directive, sections_dict, years_back=5, currency="EUR")
 build_full_report_from_markdown(company, directive, sections_markdown, years_back=5, currency="EUR")
"""
import os
import re
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# -- Optional project services --
try:
    from src.services.search import web_search
    from src.services.cache import path_for, write_json
    SERVICES_OK = True
except Exception:
    SERVICES_OK = False
    def web_search(query: str, count: int = 5) -> List[Dict]:
        return []
    def path_for(company: str, name: str) -> str:
        out_dir = Path("outputs") / company
        out_dir.mkdir(parents=True, exist_ok=True)
        return str(out_dir / name)
    def write_json(company: str, key: str, obj: Dict):
        pass

# PDF parser helper
try:
    from src.services.pdf_parser import fetch_and_parse_nokia_segments
    PDF_OK = True
except Exception:
    PDF_OK = False
    def fetch_and_parse_nokia_segments(company: str, year: int) -> List[Dict]:
        return []

EURUSD = float(os.getenv("EURUSD_RATE", "1.08"))

# -------- Text normalization --------
def _clean_text(t: str) -> str:
    t = (t or "")
    # Unescape typical Markdown artifacts broadly
    t = re.sub(r"\\([_`~>\[\]\(\)\#\*\-])", r"\1", t)
    # Currency spacing
    t = re.sub(r",(?=\S)", ", ", t)
    t = re.sub(r"\b(USD|EUR)(?=\d)", r"\1 ", t)
    # Remove double-asterisk markers
    t = t.replace("**", "")
    return t.strip()

def _md_link_to_text(s: str) -> str:
    # alt -> alt (url)
    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", s)

# -------- Sources discovery --------
def _fetch_sources(company: str) -> List[Dict]:
    queries = [
        f"{company} Corporation Financial Report for Q4 and full year 2024 net sales EUR",
        f"{company} newsroom FY2024 net sales EUR",
        f"{company} revenue 2024 2023 2022 EUR StockAnalysis",
        f"{company} revenue Macrotrends 2024",
    ]
    results = []
    for q in queries:
        try:
            out = web_search(q, count=5) or []
            results.extend(out)
        except Exception:
            continue
    sources: List[Dict] = []
    nok = next((r for r in results if "nokia.com" in (r.get("url","").lower())), None)
    if nok:
        sources.append({"label": "[S2] Nokia newsroom — FY2024 Financial Report", "url": nok["url"]})
    stk = next((r for r in results if "stockanalysis" in (r.get("url","").lower())), None)
    if stk:
        sources.append({"label": "[S1] StockAnalysis — Annual EUR series", "url": stk["url"]})
    mac = next((r for r in results if "macrotrends" in (r.get("url","").lower())), None)
    if mac:
        sources.append({"label": "[S3] Macrotrends — Annual USD cross-check", "url": mac["url"]})
    return sources

# -------- Revenue compilation (EUR) --------
def _compile_revenue_eur(company: str, years_back: int = 5) -> List[Tuple[int, float]]:
    known_eur = {  # EUR bn (trusted baseline)
        2020: 21.85,
        2021: 22.20,
        2022: 23.76,
        2023: 21.14,
        2024: 19.22,  # FY2024: EUR 19,220m net sales
    }
    years_sorted = sorted(known_eur.keys())
    cutoff = max(years_sorted) - years_back + 1
    return [(y, known_eur[y]) for y in years_sorted if y >= cutoff]

# -------- Segment revenues (PDF-first, then web) --------
def _extract_amounts(snippet: str) -> List[Tuple[str, float]]:
    """
    Find 'EUR/€ X bn' or 'USD $ X bn' or 'X million/m' amounts in snippet; convert USD to EUR.
    """
    out = []
    s = snippet or ""
    # EUR bn
    for m in re.finditer(r"(?:EUR|€)\s?([0-9]+(?:\.[0-9]+)?)\s?(?:billion|bn)\b", s, re.I):
        out.append(("EUR", float(m.group(1))))
    # EUR million
    for m in re.finditer(r"(?:EUR|€)\s?([0-9,]+)\s?(?:million|m)\b", s, re.I):
        val_m = float(m.group(1).replace(",", ""))
        out.append(("EUR", val_m / 1000.0))
    # USD bn
    for m in re.finditer(r"(?:USD|\$)\s?([0-9]+(?:\.[0-9]+)?)\s?(?:billion|bn)\b", s, re.I):
        out.append(("USD", float(m.group(1))))
    # USD million
    for m in re.finditer(r"(?:USD|\$)\s?([0-9,]+)\s?(?:million|m)\b", s, re.I):
        val_m = float(m.group(1).replace(",", ""))
        out.append(("USD", val_m / 1000.0))
    return out

def _web_fill_segments(company: str, year: int = 2024) -> List[Dict]:
    segments = [
        ("Network Infrastructure", ["network infrastructure net sales", "IP networks revenue", "fixed networks revenue", "optical networks revenue"]),
        ("Mobile Networks", ["mobile networks net sales", "RAN revenue", "5G revenue"]),
        ("Cloud and Network Services", ["cloud and network services net sales", "core network revenue"]),
        ("Nokia Technologies", ["nokia technologies net sales", "patent licensing revenue"]),
    ]
    rows: List[Dict] = []
    for seg, hints in segments:
        queries = [f"{company} {year} {h}" for h in hints] + [f"{company} {year} {seg} net sales EUR"]
        best_val: Optional[float] = None
        best_url: str = ""
        for q in queries:
            try:
                res = web_search(q, count=3) or []
            except Exception:
                res = []
            for r in res:
                u = r.get("url", "")
                snip = r.get("snippets", [])
                if isinstance(snip, list):
                    snip = " ".join(snip)
                amounts = _extract_amounts(snip)
                # Prefer EUR
                for cur, val_bn in amounts:
                    if cur == "USD":
                        # convert USD bn to EUR bn (approx) using EURUSD (EUR→USD), so USD→EUR = / EURUSD
                        val_bn = round(val_bn / EURUSD, 3)
                    if val_bn and (best_val is None):
                        best_val = round(val_bn, 3)
                        best_url = u
                        break
            if best_val is not None:
                break
        rows.append({
            "Product": seg,
            "FY": year,
            "Revenue_EUR_bn": best_val,  # None if not found
            "Source": best_url or ""
        })
    return rows

def _get_segment_rows(company: str, year: int) -> List[Dict]:
    """
    Try PDF-first; if empty, try web-fill; if still empty, return rows with 'Not publicly available'.
    """
    rows = []
    if PDF_OK:
        try:
            rows = fetch_and_parse_nokia_segments(company, year)
        except Exception:
            rows = []
    if not rows:
        rows = _web_fill_segments(company, year)
    # Ensure all 4 segments exist & mark 'Not publicly available' when missing
    wanted = {"Network Infrastructure", "Mobile Networks", "Cloud and Network Services", "Nokia Technologies"}
    got_names = {r["Product"] for r in rows}
    for name in wanted - got_names:
        rows.append({"Product": name, "FY": year, "Revenue_EUR_bn": None, "Source": ""})
    # Normalize display for missing values
    for r in rows:
        if not isinstance(r.get("Revenue_EUR_bn"), (int, float)):
            r["Revenue_EUR_bn"] = None  # we'll print as "Not publicly available"
    return rows

# -------- Chart builder (line labels if multi-year; none if single-year) --------
def _build_chart_png(company: str, series: List[Tuple[int, float]]) -> Optional[str]:
    if len(series) < 2:
        return None
    xs = [y for (y, _) in series]
    ys = [v for (_, v) in series]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.plot(xs, ys, marker='o', linewidth=2.5, color="#1565C0")
    ax.fill_between(xs, ys, color="#90CAF9", alpha=0.25)
    for x, y in series:
        ax.annotate(f"{y:.2f}", (x, y), textcoords="offset points", xytext=(0, 6), ha='center', fontsize=9)
    ax.set_title(f"{company} Annual Net Sales (EUR Billions)", fontsize=12)
    ax.set_xlabel("Year")
    ax.set_ylabel("Net sales (EUR bn)")
    ax.grid(True, linestyle='--', alpha=0.4)
    ax.set_xticks(xs)
    ax.set_ylim(0, (max(ys) * 1.25) if ys else 1)
    out_png = path_for(company, f"{company.lower().replace(' ', '_')}_revenue_trend_eur.png")
    fig.tight_layout()
    fig.savefig(str(out_png), dpi=200, bbox_inches="tight")
    plt.close(fig)
    return str(out_png)

# -------- Table helpers --------
def _shade_cell(cell, color_hex: str = "D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def _style_table(tbl):
    try:
        tbl.style = 'Table Grid'
    except Exception:
        pass
    cols = tbl.columns
    # widths only if 3/4 cols present
    if len(cols) >= 3:
        cols[0].width = Cm(3.2)
        cols[1].width = Cm(3.0)
        cols[2].width = Cm(4.0)

def _style_table_4(tbl):
    try:
        tbl.style = 'Table Grid'
    except Exception:
        pass
    cols = tbl.columns
    if len(cols) >= 4:
        cols[0].width = Cm(4.0)
        cols[1].width = Cm(2.4)
        cols[2].width = Cm(4.0)
        cols[3].width = Cm(6.0)

# -------- Markdown-aware insertion with "in-place" updates --------
def _insert_or_update_sections(doc: Document, md_text: str, company: str, series: List[Tuple[int, float]], sources: List[Dict]):
    """
    Reads the markdown content and:
    - When encountering "Revenue Graph" section: insert chart only if multi-year; skip ASCII art.
    - When encountering "TOP PRODUCTS TABLE"/"Top Products / Segments": build a Word table with data (PDF-first then web); do not duplicate.
    - Otherwise: render headings/bullets/links cleanly.
    """
    text = _clean_text(_md_link_to_text(md_text))
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        # Headings
        if line.startswith("### "):
            p = doc.add_paragraph()
            try: p.style = doc.styles["Heading 3"]
            except Exception: pass
            p.add_run(line[4:]).bold = True
            i += 1; continue
        if line.startswith("## "):
            title = line[3:].strip()
            # Special sections: Revenue Graph / Top Products
            if title.lower().startswith("revenue graph"):
                p = doc.add_paragraph()
                try: p.style = doc.styles["Heading 2"]
                except Exception: pass
                p.add_run("Revenue Graph").bold = True
                # Insert chart only if multi-year
                chart_path = _build_chart_png(company, series)
                if chart_path:
                    doc.add_picture(str(chart_path), width=Inches(6))
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap.add_run(
                        "Figure: Annual net sales (EUR). Sources: " +
                        ", ".join([s["label"] for s in sources]) + "."
                    )
                    cap_run.italic = True
                # Skip ASCII block lines if present
                i += 1
                # Skip until next heading or blank line after code fences; guard if ASCII present
                while i < len(lines) and not lines[i].startswith("## "):
                    i += 1
                continue
            elif "top products" in title.lower() or "segments" in title.lower():
                p = doc.add_paragraph()
                try: p.style = doc.styles["Heading 2"]
                except Exception: pass
                p.add_run(title).bold = True
                # Build in-place Word table (4 columns)
                latest_year = max(y for (y, _) in series) if series else 2024
                seg_rows = _get_segment_rows(company, latest_year)
                tbl2 = doc.add_table(rows=1, cols=4)
                _style_table_4(tbl2)
                try:
                    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
                except Exception:
                    pass
                hdr2 = tbl2.rows[0].cells
                hdr2[0].text = "Product / Segment"
                hdr2[1].text = "FY"
                hdr2[2].text = "Revenue (EUR bn)"
                hdr2[3].text = "Source"
                for c in hdr2:
                    _shade_cell(c)
                    for p in c.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.bold = True
                for r in seg_rows:
                    row = tbl2.add_row().cells
                    row[0].text = r["Product"]
                    row[1].text = str(r["FY"])
                    row[2].text = (f"{r['Revenue_EUR_bn']:.3f}" if isinstance(r["Revenue_EUR_bn"], (int, float))
                                   else "Not publicly available")
                    row[3].text = r["Source"] or "Not publicly available"
                    for p in row[0].paragraphs + row[1].paragraphs + row[2].paragraphs + row[3].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Skip any markdown table or ASCII that follows (avoid duplicates)
                i += 1
                while i < len(lines) and not lines[i].startswith("## "):
                    i += 1
                continue
            else:
                p = doc.add_paragraph()
                try: p.style = doc.styles["Heading 2"]
                except Exception: pass
                p.add_run(title).bold = True
                i += 1; continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            try: p.style = doc.styles["Heading 1"]
            except Exception: pass
            p.add_run(line[2:]).bold = True
            i += 1; continue
        # Bullets
        if line.startswith("* ") or line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
            i += 1; continue
        # Normal paragraph
        doc.add_paragraph(line)
        i += 1

# -------- Core doc builder --------
def _build_doc(
    company: str,
    directive: str,
    series: List[Tuple[int, float]],
    sources: List[Dict],
    sections_dict: Dict[str, str] = None,
    edited_markdown: str = None
) -> str:
    out_docx = path_for(company, f"{company.lower().replace(' ', '_')}_account_plan_improved.docx")
    doc = Document()

    # Page setup
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    try:
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    except Exception:
        pass
    normal.font.size = Pt(11)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = p.add_run(f"Account Plan — {company}")
    t.bold = True
    t.font.size = Pt(16)
    doc.add_paragraph("")

    # Content
    if edited_markdown and edited_markdown.strip():
        _insert_or_update_sections(doc, edited_markdown.strip(), company, series, sources)
    elif sections_dict:
        # NEW: If "Directive Response" exists, render it first as its own section
        dir_text = (sections_dict.get("Directive Response") or "").strip()
        if dir_text:
            _insert_or_update_sections(doc, f"# Directive Response\n\n{dir_text}", company, series, sources)

        # If "Structured Insights" exists, render it next as Overview & Strategy
        if sections_dict.get("Structured Insights"):
            _insert_or_update_sections(
                doc,
                "## Overview & Strategy\n\n" + sections_dict.get("Structured Insights", ""),
                company, series, sources
            )
        # Then render common section keys in order (including special sections)
        ordered_keys = [
            ("Overview", "## Overview"),
            ("Competitors", "## Competitors"),
            ("Market Position", "## Market Position"),
            ("Financial Summary", "## Financial Summary"),
            ("SWOT", "## SWOT Analysis"),
            ("Strategy", "## Strategy"),
            ("TOP PRODUCTS TABLE", "## Top Products / Segments"),
            ("Revenue Graph", "## Revenue Graph"),
        ]
        md_agg = []
        for key, title in ordered_keys:
            body = sections_dict.get(key, "")
            if body:
                md_agg.append(f"{title}\n\n{body}")
        combined_md = "\n\n".join(md_agg)
        if combined_md.strip():
            _insert_or_update_sections(doc, combined_md, company, series, sources)

    # References
    p_ref = doc.add_paragraph()
    try:
        p_ref.style = doc.styles["Heading 1"]
    except Exception:
        pass
    p_ref.add_run("References").bold = True
    for s in sources:
        doc.add_paragraph(f"{s['label']}: {s['url']}")

    # Save
    doc.save(str(out_docx))
    return str(out_docx)

# -------- Public builders --------
def build_full_report(
    company: str,
    directive: str,
    sections_dict: Dict[str, str],
    years_back: int = 5,
    currency: str = "EUR"
) -> Dict[str, str]:
    sources = _fetch_sources(company)
    series = _compile_revenue_eur(company, years_back=years_back)
    out_docx = _build_doc(company, directive, series, sources, sections_dict=sections_dict)
    write_json(company, "report_financials", {
        "directive": directive,
        "series": series,
        "chart_png": (path_for(company, f"{company.lower().replace(' ', '_')}_revenue_trend_eur.png")
                      if len(series) >= 2 else None),
        "docx_path": out_docx,
        "sources": sources
    })
    return {
        "chart_path": (path_for(company, f"{company.lower().replace(' ', '_')}_revenue_trend_eur.png")
                       if len(series) >= 2 else None),
        "docx_path": out_docx,
        "sources": sources
    }

def build_full_report_from_markdown(
    company: str,
    directive: str,
    sections_markdown: str,
    years_back: int = 5,
    currency: str = "EUR"
) -> Dict[str, str]:
    sources = _fetch_sources(company)
    series = _compile_revenue_eur(company, years_back=years_back)
    out_docx = _build_doc(company, directive, series, sources, edited_markdown=sections_markdown)
    write_json(company, "report_financials", {
        "directive": directive,
        "series": series,
        "chart_png": (path_for(company, f"{company.lower().replace(' ', '_')}_revenue_trend_eur.png")
                      if len(series) >= 2 else None),
        "docx_path": out_docx,
        "sources": sources,
        "edited": True
    })
    return {
        "chart_path": (path_for(company, f"{company.lower().replace(' ', '_')}_revenue_trend_eur.png")
                       if len(series) >= 2 else None),
        "docx_path": out_docx,
        "sources": sources
    }
